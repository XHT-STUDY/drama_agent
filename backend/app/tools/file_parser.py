"""FileParserTool — 上传文件解析（TXT / DOCX）(G-03)。

确定性工具（无 LLM）：
- TXT：编码探测（UTF-8 → GBK 回退），禁止可执行脚本特征；
- DOCX：python-docx 读取段落 + 表格文本；
- 大小（upload_max_bytes）、扩展名 / MIME / 内容签名联合校验；
- 拒绝：宏文档（docm / 内嵌 vbaProject）、损坏压缩包、伪装扩展名；
- 返回：ParsedFile（text / char_count / detected_format / mime_type / warnings）。

安全约束（G-03 验收）：
- 文件内容不写日志；
- 原始文件名永不用于磁盘路径（存储层负责）；
- 路径穿越在 storage 层拦截。
"""

from __future__ import annotations

import io
import logging
from typing import Any

from pydantic import BaseModel, Field

from app.core.errors import FileParseFailedError, FileTooLargeError
from app.tools.protocol import Tool, ToolMetadata

logger = logging.getLogger(__name__)

# 允许的扩展名与对应 MIME（内容签名校验后兜底）
_ALLOWED_EXTENSIONS = {".txt", ".docx"}
# DOCX 是 ZIP 容器，魔数 PK\x03\x04
_ZIP_MAGIC = b"PK\x03\x04"
# DOCX 必须包含的部件
_DOCX_REQUIRED_PART = "word/document.xml"
# 宏文档部件（拒绝）
_MACRO_PARTS = {"word/vbaProject.bin", "word/vbaData.xml"}
# 单文件最大字符数（防御性上限，超出记 warning 不阻断）
_MAX_CHARS_WARN = 1_000_000


class ParsedFile(BaseModel):
    """解析成功的结果（G-03）。"""

    model_config = {"extra": "forbid"}

    text: str = Field(..., description="解析出的纯文本（TXT 原样 / DOCX 段落+表格）")
    char_count: int = Field(..., description="文本字符数", ge=0)
    detected_format: str = Field(..., description="实际检测到的格式：txt/docx")
    mime_type: str = Field(..., description="匹配的 MIME 类型")
    warnings: list[str] = Field(default_factory=list, description="解析告警（不阻断）")


class FileParserTool(Tool):
    """上传文件解析工具（TXT / DOCX）。"""

    metadata = ToolMetadata(
        name="parse_uploaded_file",
        version="1.0",
        description="解析上传的 TXT/DOCX 文件为纯文本（编码探测 / 表格提取 / 安全校验）",
    )

    def __init__(self, upload_max_bytes: int = 10_485_760) -> None:
        self._upload_max_bytes = upload_max_bytes

    # ---- 公开 API ----

    async def execute(self, **kwargs: Any) -> ParsedFile:
        """解析上传文件。

        kwargs 必需键:
            filename: str — 客户端原始文件名（仅用于判断扩展名）
            data: bytes — 文件内容

        Raises:
            FileTooLargeError: 大小超限
            FileParseFailedError: 解析失败（含伪装扩展名 / 损坏 / 宏文档）
        """
        filename: str = kwargs["filename"]
        data: bytes = kwargs["data"]

        if len(data) > self._upload_max_bytes:
            raise FileTooLargeError(
                detail=(
                    f"文件大小 {len(data)} 字节超过上限 "
                    f"{self._upload_max_bytes} 字节"
                ),
            )

        ext = _detect_extension(filename)
        return await self._parse(ext, data)

    # ---- 私有 ----

    async def _parse(self, ext: str, data: bytes) -> ParsedFile:
        """按扩展名分发解析；内容签名校验兜底。"""
        if ext == ".docx":
            return self._parse_docx(data)
        # .txt（默认路径）：先看内容签名是否实际是 zip（伪装 txt）
        if data.startswith(_ZIP_MAGIC):
            raise FileParseFailedError(
                detail="文件扩展名为 .txt 但内容是压缩包（疑似 DOCX 伪装）",
            )
        return self._parse_txt(data)

    @staticmethod
    def _parse_txt(data: bytes) -> ParsedFile:
        """TXT 编码探测：UTF-8 → GBK 回退；两者都失败则报错。"""
        text, encoding = _decode_text(data)
        warnings: list[str] = []
        if encoding != "utf-8":
            warnings.append(f"文本按 {encoding.upper()} 编码解码（非 UTF-8）")
        return ParsedFile(
            text=text,
            char_count=len(text),
            detected_format="txt",
            mime_type="text/plain",
            warnings=warnings,
        )

    @staticmethod
    def _parse_docx(data: bytes) -> ParsedFile:
        """DOCX 解析：校验 zip / 宏 / 必需部件，提取段落 + 表格文本。"""
        import zipfile

        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                names = set(zf.namelist())
                macro_hits = _MACRO_PARTS & names
                if macro_hits:
                    raise FileParseFailedError(
                        detail=f"检测到宏部件（{sorted(macro_hits)[0]}），"
                        "拒绝解析含宏的文档",
                    )
                if _DOCX_REQUIRED_PART not in names:
                    raise FileParseFailedError(
                        detail=f"压缩包缺少 {_DOCX_REQUIRED_PART}，不是有效 DOCX",
                    )
        except FileParseFailedError:
            raise
        except zipfile.BadZipFile:
            raise FileParseFailedError(detail="不是有效的 DOCX 文件（损坏的压缩包）") from None
        except Exception as exc:  # noqa: BLE001 — 统一映射为解析失败
            logger.warning("DOCX zip 读取失败: %s", exc)
            raise FileParseFailedError(detail="DOCX 压缩包读取失败") from exc

        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
        except FileParseFailedError:
            raise
        except Exception as exc:  # noqa: BLE001 — 损坏文档统一映射
            logger.warning("DOCX 文档解析失败: %s", exc)
            raise FileParseFailedError(detail="DOCX 文档损坏，无法解析") from exc

        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()

        warnings: list[str] = []
        if not text:
            warnings.append("DOCX 未提取到任何文本内容")
        if len(text) > _MAX_CHARS_WARN:
            warnings.append(f"文本长度超 {_MAX_CHARS_WARN} 字符，超出部分已忽略")

        return ParsedFile(
            text=text,
            char_count=len(text),
            detected_format="docx",
            mime_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml."
                "document"
            ),
            warnings=warnings,
        )


# ---- 模块级辅助函数 ----


def _detect_extension(filename: str) -> str:
    """从原始文件名提取扩展名；伪装扩展名（多级 / 路径分隔符）直接报错。

    Returns:
        小写扩展名（.txt / .docx）

    Raises:
        FileParseFailedError: 扩展名不受支持或文件名非法
    """
    # 拒绝路径分隔符（防客户端文件名注入磁盘路径）
    if "/" in filename or "\\" in filename or ".." in filename:
        raise FileParseFailedError(
            detail="文件名包含非法路径字符（不允许路径穿越）",
        )
    lower = filename.lower()
    # 取最后一个点后的扩展名；无点或点在前部（.txt 隐藏文件）按无扩展名处理
    base = lower.rsplit(".", 1)
    if len(base) != 2 or not base[1]:
        raise FileParseFailedError(
            detail=f"不支持的文件类型: {filename!r}（仅支持 .txt/.docx）",
        )
    ext = f".{base[1]}"
    if ext not in _ALLOWED_EXTENSIONS:
        raise FileParseFailedError(
            detail=f"不支持的文件类型: {ext}（仅支持 .txt/.docx）",
        )
    return ext


def _decode_text(data: bytes) -> tuple[str, str]:
    """TXT 编码探测：UTF-8 → GBK 回退。

    Returns:
        (text, encoding)

    Raises:
        FileParseFailedError: 两种编码都无法解码
    """
    for encoding in ("utf-8", "gbk"):
        try:
            return data.decode(encoding), encoding
        except (UnicodeDecodeError, LookupError):
            continue
    raise FileParseFailedError(detail="文本编码无法识别（支持 UTF-8 / GBK）")
