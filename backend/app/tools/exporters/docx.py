"""DOCX 序列化器 (G-05) — 用 python-docx 把导出 Markdown 转成可打开的 Word 文档。

满足 G-05 验收「DOCX 可打开，中文、表格和分页正常」：
- 中文字体 fallback：正文与各级标题统一设置 w:eastAsia（Word 中文环境的
  常见可用字体，系统缺省会回退）；
- 页眉：文档标题；页脚：PAGE 域页码；
- 分页：每个一级标题新起一页（首个文档标题除外）；
- 表格：支持 GFM 管道表格（维度得分等结构化数据可进 Word 表格）。

模块边界：纯序列化 + DocxExporter Tool，不触碰 API / 存储。
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.tools.protocol import Tool, ToolMetadata

# Word 中文环境常见可用字体；若系统缺省会回退到默认字体
_CN_FONT = "宋体"
_LATIN_FONT = "Calibri"

_HEADING_STYLE_NAMES = ("Heading 1", "Heading 2", "Heading 3")


def _apply_east_asia_font(style: Any) -> None:
    """给样式设置 w:eastAsia 中文字体（python-docx 默认不设置中文字体）。

    style.font.name 只设置 w:ascii / w:hAnsi（西文），中文渲染还需 eastAsia。
    """
    style.font.name = _LATIN_FONT
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), _CN_FONT)


def _add_page_number_field(paragraph: Any) -> None:
    """在段落中插入 PAGE 域（Word 自动编号）。"""
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def _add_markdown_table(doc: _Document, rows: list[list[str]]) -> None:
    """把 GFM 管道表格解析出的行转成 Word 表格（带边框，便于阅读）。"""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(n_cols):
            cell_text = row[c] if c < len(row) else ""
            table.cell(r, c).text = cell_text


def _parse_table_header(line: str) -> list[str]:
    """解析 GFM 表格表头行（去掉首尾 |，按 | 拆列并去空格）。"""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def markdown_to_docx(markdown: str, *, header_text: str = "") -> _Document:
    """把导出 Markdown 转成 python-docx Document。

    行解析规则与前端 buildDocxBlob 对齐：
    # → Heading 1（新起一页）、## → Heading 2、### → Heading 3、
    `  - ` → 二级列表、`- ` / `* ` → 一级列表、`> ` → 斜体引用、
    GFM 管道表格 → Word 表格、其余 → 普通段落。

    Args:
        markdown: 导出 Markdown 全文
        header_text: 页眉文本（通常为项目名）

    Returns:
        python-docx Document（未保存）
    """
    doc = Document()

    # 中文字体 fallback：正文 + 各级标题
    _apply_east_asia_font(doc.styles["Normal"])
    for name in _HEADING_STYLE_NAMES:
        try:
            _apply_east_asia_font(doc.styles[name])
        except KeyError:
            continue

    # 页眉（文档标题）
    if header_text:
        section = doc.sections[0]
        header_para = section.header.paragraphs[0]
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 页脚页码（PAGE 域）
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(footer_para)

    lines = markdown.split("\n")
    first_heading_seen = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # GFM 管道表格：表头行后紧跟分隔行（含 ---）
        if line.startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            header_row = _parse_table_header(line)
            i += 2  # 跳过分隔行
            rows = [header_row]
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_parse_table_header(lines[i]))
                i += 1
            _add_markdown_table(doc, rows)
            continue

        stripped = line.strip()
        if stripped == "":
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            heading = doc.add_heading(line[2:], level=1)
            if first_heading_seen:
                heading.paragraph_format.page_break_before = True
            first_heading_seen = True
        elif line.startswith("  - "):
            doc.add_paragraph(line[4:], style="List Bullet 2")
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("> "):
            para = doc.add_paragraph()
            run = para.add_run(line[2:])
            run.italic = True
        else:
            doc.add_paragraph(line)
        i += 1

    return doc


def build_docx_bytes(markdown: str, *, header_text: str = "") -> bytes:
    """生成 .docx 字节流（供 FileStore 落盘）。"""
    doc = markdown_to_docx(markdown, header_text=header_text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class DocxExporter(Tool):
    """DOCX 序列化工具（G-05）。"""

    metadata = ToolMetadata(
        name="export_docx",
        version="1.0",
        description="把导出 Markdown 转成可打开的 Word 文档（中文字体/页眉/页码/分页）",
    )

    async def execute(self, **kwargs: Any) -> dict[str, Any]:
        """把 Markdown 转成 .docx 字节。

        Args:
            markdown: 导出 Markdown 全文
            header_text: 可选页眉文本（项目名）

        Returns:
            {"data": bytes, "size_bytes": int}
        """
        markdown: str = kwargs["markdown"]
        header_text: str = kwargs.get("header_text", "")
        data = build_docx_bytes(markdown, header_text=header_text)
        return {"data": data, "size_bytes": len(data)}
