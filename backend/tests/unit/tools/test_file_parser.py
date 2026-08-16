"""FileParserTool 单元测试 (G-03)。

覆盖：
- TXT：UTF-8 直接解码、GBK 回退（含 warning）、空文本、编码无法识别拒绝;
- DOCX：段落 + 表格文本提取、中文内容、表格以 ` | ` 连接;
- 大小超限（FileTooLargeError）;
- 拒绝场景：不支持扩展名、无扩展名、路径穿越文件名、.txt 伪装 zip、
  DOCX 非 zip、DOCX 缺必需部件、DOCX 含宏部件;
- 返回结构：detected_format / mime_type / char_count / warnings。
"""

from __future__ import annotations

import io
import uuid
import zipfile

import pytest
from pydantic import ValidationError

from app.core.errors import FileParseFailedError, FileTooLargeError
from app.tools.file_parser import FileParserTool, ParsedFile


def _build_docx(
    paragraphs: list[str] | None = None, rows: list[list[str]] | None = None
) -> bytes:
    """用 python-docx 在内存中构造一个真实 DOCX 字节。"""
    from docx import Document

    doc = Document()
    for p in paragraphs or []:
        doc.add_paragraph(p)
    for row in rows or []:
        table = doc.add_table(rows=1, cols=len(row))
        for i, cell in enumerate(row):
            table.rows[0].cells[i].text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _run(
    filename: str, data: bytes, *, max_bytes: int = 10_485_760
) -> ParsedFile:
    return await FileParserTool(upload_max_bytes=max_bytes).execute(
        filename=filename, data=data,
    )


@pytest.mark.asyncio
class TestTxtParsing:
    """TXT 解析：编码探测与告警。"""

    async def test_utf8_txt(self) -> None:
        """UTF-8 中文文本直接解码，无 warning。"""
        data = "第一集：夜城初遇".encode()
        parsed = await _run("story.txt", data)
        assert parsed.detected_format == "txt"
        assert parsed.mime_type == "text/plain"
        assert parsed.text == "第一集：夜城初遇"
        assert parsed.char_count == len("第一集：夜城初遇")
        assert parsed.warnings == []

    async def test_gbk_txt_fallback(self) -> None:
        """GBK 编码回退成功，并产生编码告警。"""
        text = "夜城的霓虹在雨里融化"
        data = text.encode("gbk")
        parsed = await _run("剧.txt", data)
        assert parsed.detected_format == "txt"
        assert parsed.text == text
        assert any("GBK" in w for w in parsed.warnings)

    async def test_utf8_bytes_look_like_gbk_but_valid_utf8(self) -> None:
        """UTF-8 字节同时可能被 GBK 解码时，优先 UTF-8。"""
        text = "短剧行业标准第一稿"
        data = text.encode("utf-8")
        parsed = await _run("x.txt", data)
        assert parsed.text == text
        assert parsed.warnings == []  # 按 UTF-8 解码，无告警

    async def test_empty_txt(self) -> None:
        """空文件解析为空文本，char_count=0。"""
        parsed = await _run("empty.txt", b"")
        assert parsed.text == ""
        assert parsed.char_count == 0

    async def test_undecodable_bytes(self) -> None:
        """既非 UTF-8 也非 GBK 的字节 → FileParseFailedError。"""
        data = b"\xff\xfe\x00\x11\x22"  # 无效 UTF-8 / GBK 序列
        with pytest.raises(FileParseFailedError):
            await _run("bad.txt", data)


@pytest.mark.asyncio
class TestDocxParsing:
    """DOCX 解析：段落 + 表格提取。"""

    async def test_paragraphs_and_tables(self) -> None:
        """段落与表格文本都被提取，表格单元格以 ` | ` 连接。"""
        data = _build_docx(
            paragraphs=["第一集", "夜城初遇"],
            rows=[["角色", "台词"], ["林晚", "你是谁"]],
        )
        parsed = await _run("script.docx", data)
        assert parsed.detected_format == "docx"
        assert "第一集" in parsed.text
        assert "夜城初遇" in parsed.text
        assert "角色 | 台词" in parsed.text
        assert "林晚 | 你是谁" in parsed.text
        assert parsed.char_count == len(parsed.text)

    async def test_chinese_docx_encoding(self) -> None:
        """中文 DOCX 提取不乱码。"""
        data = _build_docx(paragraphs=["霓虹闪烁的夜城", "他推门而入"])
        parsed = await _run("短剧.docx", data)
        assert parsed.text == "霓虹闪烁的夜城\n他推门而入"

    async def test_empty_docx_warns(self) -> None:
        """无文本 DOCX 产生告警但不失败。"""
        data = _build_docx(paragraphs=[""])  # 仅空段落
        parsed = await _run("empty.docx", data)
        assert any("未提取到任何文本" in w for w in parsed.warnings)


@pytest.mark.asyncio
class TestRejections:
    """拒绝场景：大小、扩展名、伪装、损坏、宏、路径穿越。"""

    async def test_too_large(self) -> None:
        """超过 upload_max_bytes → FileTooLargeError。"""
        with pytest.raises(FileTooLargeError):
            await _run("big.txt", b"a" * 2048, max_bytes=1024)

    async def test_unsupported_extension(self) -> None:
        """不支持扩展名 → FileParseFailedError。"""
        for name in ("a.pdf", "b.md", "c.exe", "d.bin"):
            with pytest.raises(FileParseFailedError):
                await _run(name, b"hello")

    async def test_missing_extension(self) -> None:
        """无扩展名 / 隐藏文件式点名 → FileParseFailedError。"""
        with pytest.raises(FileParseFailedError):
            await _run("README", b"hello")
        with pytest.raises(FileParseFailedError):
            await _run(".gitignore", b"hello")

    async def test_uppercase_extension_accepted(self) -> None:
        """大小写不敏感：.TXT / .DocX 均可。"""
        parsed = await _run("A.TXT", "内容".encode())
        assert parsed.detected_format == "txt"
        parsed2 = await _run("B.DocX", _build_docx(paragraphs=["x"]))
        assert parsed2.detected_format == "docx"

    async def test_path_traversal_filename(self) -> None:
        """文件名含路径分隔符 / `..` → 拒绝。"""
        for name in ("../etc/passwd.txt", "sub/../a.txt", "..\\win.txt"):
            with pytest.raises(FileParseFailedError):
                await _run(name, b"hello")

    async def test_txt_disguising_zip(self) -> None:
        """.txt 文件名但内容是 zip（伪装 DOCX）→ 拒绝。"""
        real_docx = _build_docx(paragraphs=["假装是文本"])
        with pytest.raises(FileParseFailedError):
            await _run("sneaky.txt", real_docx)

    async def test_docx_not_a_zip(self) -> None:
        """.docx 但内容不是 zip → FileParseFailedError。"""
        with pytest.raises(FileParseFailedError):
            await _run("fake.docx", b"this is not a zip")

    async def test_docx_missing_required_part(self) -> None:
        """zip 但缺 word/document.xml → FileParseFailedError。"""
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("hello.txt", "hi")
        with pytest.raises(FileParseFailedError):
            await _run("broken.docx", buf.getvalue())

    async def test_docx_with_macro_part(self) -> None:
        """含宏部件（vbaProject）→ FileParseFailedError。"""
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("word/document.xml", "<doc/>")
            zf.writestr("word/vbaProject.bin", b"\x00MACRO")
        with pytest.raises(FileParseFailedError):
            await _run("macro.docx", buf.getvalue())


@pytest.mark.asyncio
class TestResultShape:
    """返回结构校验。"""

    async def test_parsed_file_extra_forbidden(self) -> None:
        """ParsedFile 拒绝未知字段（extra=forbid）。"""
        from app.tools.file_parser import ParsedFile

        with pytest.raises(ValidationError):
            ParsedFile.model_validate(
                {
                    "text": "x",
                    "char_count": 1,
                    "detected_format": "txt",
                    "mime_type": "text/plain",
                    "bogus": 1,
                }
            )

    async def test_return_metadata(self) -> None:
        """返回结构包含 sha256 所需的全部字段。"""
        data = "正文内容".encode()
        parsed = await _run("meta.txt", data)
        assert parsed.text == "正文内容"
        assert parsed.char_count == 4
        assert parsed.detected_format == "txt"
        assert parsed.mime_type == "text/plain"
        assert parsed.warnings == []

    async def test_random_key_not_used_for_parser(self) -> None:
        """解析器不关心文件名是否服务端随机——只取扩展名。"""
        name = f"{uuid.uuid4().hex}.txt"
        parsed = await _run(name, "随机存储键文件名".encode())
        assert parsed.detected_format == "txt"
