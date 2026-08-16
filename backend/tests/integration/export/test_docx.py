"""DOCX 导出集成测试 (G-05)。

用 python-docx 重新打开生成的 .docx，验证 G-05 验收项：
- 中文不乱码（标题 / 正文 / 页眉）;
- 表格正常（GFM 管道表格 → Word 表格）;
- 分页正常（除文档抬头外每个一级标题新起一页）;
- 页码域存在（页脚 PAGE 字段）;
- DocxExporter Tool 返回可被 FileStore 落盘的字节流。

纯序列化测试，不触碰 DB。
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document

from app.tools.exporters.docx import DocxExporter, build_docx_bytes, markdown_to_docx
from app.tools.exporters.markdown import build_export_markdown


def _sample_markdown() -> str:
    """构造包含标题 / 中文正文 / 引用 / 列表 / 表格的导出 Markdown。"""
    return "\n".join(
        [
            "# 足球少年逆袭记 — 内容导出",
            "",
            "> 导出时间：2026-08-16T10:00:00+00:00",
            "",
            "## 基本信息",
            "",
            "- 剧名：足球少年逆袭记",
            "- 类型：热血运动",
            "",
            "## 维度得分",
            "",
            "| 维度 | 得分 |",
            "| --- | --- |",
            "| 开头钩子 | 90 分 |",
            "| 节奏控制 | 78 分 |",
            "",
            "# 十集大纲",
            "",
            "第 1 集：被抛弃",
            "",
        ]
    )


def _reopen_document(data: bytes) -> Any:
    """用 python-docx 重新打开生成的字节流（模拟用户打开 .docx）。

    返回 Any：python-docx 的 Document 是工厂函数，无法用作 mypy 类型。
    """
    return Document(BytesIO(data))


class TestDocxContent:
    """中文与内容结构。"""

    def test_chinese_preserved_in_paragraphs(self) -> None:
        """中文标题与正文不丢失、不乱码。"""
        doc = markdown_to_docx(_sample_markdown(), header_text="足球少年逆袭记")
        texts = [p.text for p in doc.paragraphs]
        assert any("内容导出" in t for t in texts), "文档标题未进入正文"
        assert any("剧名：足球少年逆袭记" in t for t in texts), "中文正文丢失"
        assert any("导出时间：2026-08-16" in t for t in texts), "块引用文本丢失"

    def test_heading_styles(self) -> None:
        """一级 / 二级标题使用 Word 标题样式。"""
        doc = markdown_to_docx(_sample_markdown())
        styles = [p.style.name if p.style else "" for p in doc.paragraphs if p.text.strip()]
        assert "Heading 1" in styles
        assert "Heading 2" in styles
        assert styles[0] == "Heading 1", "文档抬头应为 Heading 1"

    def test_table_created(self) -> None:
        """GFM 管道表格 → Word 表格（验收项：表格正常）。"""
        doc = markdown_to_docx(_sample_markdown())
        assert len(doc.tables) == 1, "应生成一个表格"
        table = doc.tables[0]
        assert table.cell(0, 0).text == "维度"
        assert table.cell(0, 1).text == "得分"
        assert table.cell(1, 0).text == "开头钩子"
        assert table.cell(2, 1).text == "78 分"

    def test_header_text(self) -> None:
        """页眉写入项目名。"""
        doc = markdown_to_docx(_sample_markdown(), header_text="足球少年逆袭记")
        header_text = doc.sections[0].header.paragraphs[0].text
        assert header_text == "足球少年逆袭记"


class TestDocxPagination:
    """分页与页码。"""

    def test_h1_page_break_except_first(self) -> None:
        """首个标题不强制分页；其后的一级标题均 page_break_before。"""
        doc = markdown_to_docx(_sample_markdown())
        h1s = [
            p
            for p in doc.paragraphs
            if (p.style.name if p.style else "") == "Heading 1" and p.text.strip()
        ]
        assert len(h1s) >= 2
        # 第一个（文档抬头）不分页
        assert h1s[0].paragraph_format.page_break_before is not True
        # 其后的一级标题（十集大纲）新起一页
        for h in h1s[1:]:
            assert h.paragraph_format.page_break_before is True, f"{h.text} 未分页"

    def test_page_number_field_in_footer(self) -> None:
        """页脚包含 PAGE 域（Word 自动页码）。"""
        doc = markdown_to_docx(_sample_markdown())
        footer_xml = doc.sections[0].footer._element.xml
        assert "PAGE" in footer_xml, "页脚应包含 PAGE 域"


class TestDocxExporterTool:
    """DocxExporter Tool 产出可重开的字节流。"""

    async def test_execute_returns_openable_bytes(self) -> None:
        """execute 返回 data / size_bytes，且 data 可被 python-docx 重开。"""
        exporter = DocxExporter()
        result = await exporter.execute(
            markdown=_sample_markdown(), header_text="足球少年逆袭记"
        )
        assert result["size_bytes"] == len(result["data"])
        doc = _reopen_document(result["data"])
        assert any("内容导出" in p.text for p in doc.paragraphs)

    async def test_full_export_pipeline_to_docx(self) -> None:
        """build_export_markdown → build_docx_bytes 全链路（含中文）。"""
        from tests.unit.export.test_markdown import _data

        md = build_export_markdown(
            project_title="足球少年逆袭记",
            exported_at="2026-08-16T10:00:00+00:00",
            data=_data(),
            kinds=["story_bible", "outline", "script", "evaluation", "revision"],
        )
        data: bytes = build_docx_bytes(md, header_text="足球少年逆袭记")
        doc = _reopen_document(data)
        texts = "".join(p.text for p in doc.paragraphs)
        assert "世界观与人物设定" in texts
        assert "十集大纲" in texts
        assert "第 1 集剧本：被抛弃" in texts
        assert "评估报告" in texts
        assert "修订说明" in texts
